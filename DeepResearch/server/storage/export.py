import os
from typing import Optional
from docx import Document

from config import settings
from storage.pdf import PDFExporter
from utils.logger import logger
from utils.markdown import markdown_to_html


class ReportExportService:
    """
    Unified Report Export Service converting Markdown reports
    to PDF, DOCX, HTML, and Markdown export files.
    """

    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or os.path.join(settings.UPLOAD_DIR, "reports")
        os.makedirs(self.output_dir, exist_ok=True)
        self.pdf_exporter = PDFExporter(output_dir=self.output_dir)

    def export(self, title: str, markdown_content: str, format_type: str = "pdf") -> Optional[str]:
        """
        Export markdown report to specified target format.
        """
        fmt = format_type.lower()
        safe_title = "".join([c if c.isalnum() else "_" for c in title])

        if fmt == "pdf":
            html = markdown_to_html(markdown_content)
            return self.pdf_exporter.generate_pdf(html, f"{safe_title}.pdf")
        elif fmt == "docx":
            return self._export_docx(title, markdown_content, f"{safe_title}.docx")
        elif fmt == "html":
            return self._export_html(title, markdown_content, f"{safe_title}.html")
        elif fmt in ["md", "markdown"]:
            return self._export_markdown(markdown_content, f"{safe_title}.md")
        else:
            logger.error(f"[ReportExportService] Unsupported format: {format_type}")
            return None

    def _export_docx(self, title: str, markdown_content: str, filename: str) -> str:
        filepath = os.path.join(self.output_dir, filename)
        doc = Document()
        doc.add_heading(title, 0)

        for paragraph in markdown_content.split("\n\n"):
            p = paragraph.strip()
            if p.startswith("# "):
                doc.add_heading(p.replace("# ", ""), level=1)
            elif p.startswith("## "):
                doc.add_heading(p.replace("## ", ""), level=2)
            elif p.startswith("### "):
                doc.add_heading(p.replace("### ", ""), level=3)
            elif p:
                doc.add_paragraph(p)

        doc.save(filepath)
        logger.info(f"[ReportExportService] DOCX successfully generated at {filepath}")
        return filepath

    def _export_html(self, title: str, markdown_content: str, filename: str) -> str:
        filepath = os.path.join(self.output_dir, filename)
        html_body = markdown_to_html(markdown_content)
        full_html = f"<!DOCTYPE html><html><head><title>{title}</title></head><body>{html_body}</body></html>"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(full_html)
        return filepath

    def _export_markdown(self, markdown_content: str, filename: str) -> str:
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        return filepath
